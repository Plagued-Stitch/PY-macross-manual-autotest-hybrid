
import os

from docx import Document
from docx.shared import Inches
from mss import mss

class BugReport:
    #В случае фейла PyTest запускает метод который создает дирректорию в документах пользователя, создаёт в ней скриншоты и баг-репорты
    def write(testCase='', caseStep='', expectResult='', actualResult=''):
        documentsFolder = F'{os.getenv("USERPROFILE")}\\Documents\\1C Tests Bug-reports\\'
        document = Document()

        if not os.path.exists(documentsFolder):
            os.makedirs(documentsFolder)

        with mss() as sct:
            sct.shot(output=F'{documentsFolder}bug-{testCase}-{caseStep}.png')

        table = document.add_table(rows=5, cols=2, style='TableGrid')

        table.cell(0, 0).add_paragraph(text='Тест-кейс:')
        table.cell(1, 0).add_paragraph(text='Шаг тест-кейса:')
        table.cell(2, 0).add_paragraph(text='Ожидаемый результат:')
        table.cell(3, 0).add_paragraph(text='Фактический результат:')
        table.cell(4, 0).add_paragraph(text='Окружение:')

        table.cell(0, 1).add_paragraph().add_run(testCase).bold = True
        table.cell(1, 1).add_paragraph().add_run(caseStep).bold = True
        table.cell(2, 1).add_paragraph().add_run(expectResult).bold = True
        table.cell(3, 1).add_paragraph().add_run(actualResult).bold = True
        table.cell(4, 1).add_paragraph().add_run('Windows 10 Pro - 1909 (18363.1556); 1C V8 Training 8.3.24.1342').bold = True

        document.add_picture(F'{documentsFolder}bug-{testCase}-{caseStep}.png', width=Inches(7), height=Inches(4))

        return document.save(F'{documentsFolder}bugreport-{testCase}-{caseStep}.docx')